from mcp.server.fastmcp import FastMCP
from pydantic import BaseModel, Field
from typing import Optional
import docx
import json

mcp = FastMCP("docx_mcp")

class DocxFilePath(BaseModel):
    file_path: str = Field(..., description="Path to the DOCX file (e.g., 'C:/data/file.docx')")

class CreateDocxInput(BaseModel):
    file_path: str = Field(..., description="Path to save the new DOCX file.")
    content: str = Field(..., description="The text content to add to the document.")

@mcp.tool(
    name="docx_extract_text",
    annotations={
        "readOnlyHint": True,
    }
)
async def extract_text(params: DocxFilePath) -> str:
    """Extracts all text from a DOCX file."""
    try:
        doc = docx.Document(params.file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append("".join(run.text for run in para.runs))
        return '\n'.join(full_text)
    except Exception as e:
        return f"Error extracting text from {params.file_path}: {e}"

@mcp.tool(
    name="docx_extract_tables",
    annotations={
        "readOnlyHint": True,
    }
)
async def extract_tables(params: DocxFilePath) -> str:
    """Extracts all tables from a DOCX file."""
    try:
        doc = docx.Document(params.file_path)
        tables_data = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = "".join(para.text for para in cell.paragraphs)
                    row_data.append(cell_text)
                table_data.append(row_data)
            tables_data.append(table_data)
        return json.dumps(tables_data, indent=2)
    except Exception as e:
        return f"Error extracting tables from {params.file_path}: {e}"

@mcp.tool(
    name="docx_create_document",
    annotations={
        "destructiveHint": False, # Not destructive as it creates a new file
    }
)
async def create_document(params: CreateDocxInput) -> str:
    """Creates a new DOCX file with the given content."""
    try:
        doc = docx.Document()
        doc.add_paragraph(params.content)
        doc.save(params.file_path)
        return f"Successfully created DOCX file at {params.file_path}"
    except Exception as e:
        return f"Error creating DOCX file at {params.file_path}: {e}"

if __name__ == "__main__":
    mcp.run()
